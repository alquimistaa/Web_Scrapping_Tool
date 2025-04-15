import os
import requests
from bs4 import BeautifulSoup
from bs4.element import Tag
from urllib.parse import urljoin, urlparse
from docx import Document

# Gereksiz HTML etiketleri ve blacklist anahtar kelimeler
REMOVABLE_TAGS = ["script", "style", "nav", "footer", "header", "aside"]
BLACKLIST_KEYWORDS = [
    "instagram", "linkedin", "youtube", "twitter", "facebook",
    "copyright", "çerez", "©", "tümünü oku", "tamam",
    "gizlilik politikası", "çerez politikası"
]


def scrape_site(base_url: str):
    parsed_url = urlparse(base_url)
    domain = parsed_url.netloc
    output_dir = f"output_{domain.replace('.', '_')}"
    os.makedirs(output_dir, exist_ok=True)

    # Sayfa linklerini topla
    try:
        response = requests.get(base_url)
        response.encoding = response.apparent_encoding
        soup = BeautifulSoup(response.text, "html.parser")
    except Exception as e:
        print(f"🔴 Ana sayfa alınamadı: {e}")
        return

    all_links = set()
    for a_tag in soup.find_all("a", href=True):
        full_url = urljoin(base_url, a_tag['href'])
        if parsed_url.netloc in urlparse(full_url).netloc:
            all_links.add(full_url)
    all_links.add(base_url)  # Ana sayfa da dahil

    # Her sayfa için scraping işlemi
    for page_url in all_links:
        try:
            res = requests.get(page_url)
            html_content = res.content.decode("utf-8", errors="replace")
            page_soup = BeautifulSoup(html_content, "html.parser")

            for tag in page_soup.find_all(True):
                if not isinstance(tag, Tag):
                    continue
                try:
                    tag_name = tag.name.lower()
                    tag_classes = [cls.lower() for cls in tag.get("class", []) if isinstance(cls, str)]
                    tag_id = tag.get("id", "").lower() if isinstance(tag.get("id"), str) else ""
                    if tag_name in REMOVABLE_TAGS or any("menu" in cls for cls in tag_classes) or "menu" in tag_id:
                        tag.decompose()
                except Exception:
                    continue

            for a in page_soup.find_all("a"):
                a.unwrap()

            elements = page_soup.find_all(["p", "h1", "h2", "h3", "h4", "h5"])
            unique_texts = []
            for el in elements:
                text = ''.join(el.strings).strip()
                if not text:
                    continue
                lower_text = text.lower()
                if any(keyword in lower_text for keyword in BLACKLIST_KEYWORDS):
                    continue
                if text not in unique_texts:
                    unique_texts.append(text)

            # Dosya adını oluştur
            path = urlparse(page_url).path.strip("/") or "anasayfa"
            filename = f"{path.replace('/', '_')}.docx"
            filepath = os.path.join(output_dir, filename)

            doc = Document()
            for line in unique_texts:
                if line.isupper() or line.endswith("!") or len(line.split()) < 5:
                    doc.add_paragraph(line, style="Heading 2")
                else:
                    doc.add_paragraph(line)

            doc.save(filepath)
            print(f"✅ {filename} kaydedildi.")
        except Exception as e:
            print(f"❌ {page_url} hatalı: {e}")

    print(f"\n📁 Tüm dosyalar klasöre kaydedildi: {output_dir}")
